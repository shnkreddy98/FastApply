from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from datetime import datetime
from typing import Dict, List, Any
from dataclasses import dataclass
from docx2pdf import convert

@dataclass
class Section:
    title: str
    content_generator: callable

class ResumeGenerator:
    MARGIN = 0.5
    DATE_INPUT_FORMAT = '%Y-%m-%d'
    DATE_OUTPUT_FORMAT = '%b %Y'
    NAME_FONT_SIZE = 14
    DEFAULT_FONT_SIZE = 11

    def __init__(self, template_path: str):
        self.template = Document(template_path)
        self.document = Document()
        self._setup_margins()

    def _setup_margins(self) -> None:
        for section in self.document.sections:
            for margin in ['top_margin', 'bottom_margin', 'left_margin', 'right_margin']:
                setattr(section, margin, Inches(self.MARGIN))

    def _format_date(self, date_str: str) -> str:
        try:
            date_obj = datetime.strptime(date_str, self.DATE_INPUT_FORMAT)
            return date_obj.strftime(self.DATE_OUTPUT_FORMAT).capitalize()
        except ValueError as e:
            raise ValueError(f"Invalid date format. Expected YYYY-MM-DD, got {date_str}") from e

    def _add_bottom_border(self, paragraph) -> None:
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        bottom = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="auto"/></w:pBdr>')
        pPr.append(bottom)

    def _add_paragraph(self, text: str = "", bold: bool = False, spacing: int = 1, 
                      font_size: int = DEFAULT_FONT_SIZE, align_center: bool = False,
                      add_border: bool = False) -> None:
        paragraph = self.document.add_paragraph()
        if text:
            run = paragraph.add_run(text)
            run.bold = bold
            font = run.font
            font.size = Pt(font_size)
            font.name = "Times New Roman" 
        if align_center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(spacing)
        paragraph.paragraph_format.space_after = Pt(spacing)
        if add_border:
            self._add_bottom_border(paragraph)
        return paragraph

    def _add_two_column_table(self, left_text: str, right_text: str, bold: bool = False) -> None:
        table = self.document.add_table(rows=1, cols=2)
        table.autofit = True
        table.allow_autofit = True
        left_cell = table.cell(0, 0)
        left_paragraph = left_cell.paragraphs[0]
        left_run = left_paragraph.add_run(left_text)
        left_run.bold = bold
        left_run.font.size = Pt(self.DEFAULT_FONT_SIZE)
        left_run.font.name = "Times New Roman"
        left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_paragraph.paragraph_format.space_before = Pt(1)
        left_paragraph.paragraph_format.space_after = Pt(1)

        right_cell = table.cell(0, 1)
        right_paragraph = right_cell.paragraphs[0]
        right_run = right_paragraph.add_run(right_text)
        right_run.font.size = Pt(self.DEFAULT_FONT_SIZE)
        right_run.font.name = "Times New Roman"
        right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_paragraph.paragraph_format.space_before = Pt(1)
        right_paragraph.paragraph_format.space_after = Pt(1)

    def _add_header_section(self, data: Dict[str, Any]) -> None:
        self._add_paragraph(data['user_name'], bold=True, font_size=self.NAME_FONT_SIZE, align_center=True)
        contact_info = f"{data['user_city']} | {data['user_phone']} | {data['user_email']}"
        if data['user_linkedin_link']:
            contact_info += f" | {data['user_linkedin_link']}"
        self._add_paragraph(contact_info, align_center=True)
        self._add_paragraph("Summary", bold=True, add_border=True)
        self._add_paragraph(data['user_summary'])

    def _add_experience_section(self, experiences: List[Dict[str, Any]]) -> None:
        self._add_paragraph("Work Experience", bold=True, add_border=True)
        for exp in experiences:
            left_text = f"{exp['exp_role']} | {exp['exp_company']}"
            right_text = f"{self._format_date(exp['exp_start_date'])} - {self._format_date(exp['exp_end_date'])}"
            self._add_two_column_table(left_text, right_text, bold=True)
            description_points = exp['exp_description'].split('.')
            for point in description_points:
                if point.strip():
                    self._add_paragraph(f"• {point.strip()}", spacing=1)

    def _add_skills_section(self, skills: str) -> None:
        self._add_paragraph("Skills", bold=True, add_border=True)
        for skill in skills.split('\n'):
            if skill.strip():
                self._add_paragraph(f"{skill.strip()}.", spacing=1)

    def _add_projects_section(self, projects: List[Dict[str, Any]]) -> None:
        self._add_paragraph("Academic Projects", bold=True, add_border=True)
        for project in projects:
            self._add_two_column_table(
                project['proj_name'],
                f"{self._format_date(project['proj_start_date'])} - {self._format_date(project['proj_end_date'])}",
                bold=True
            )
            description_points = project['proj_description'].split('.')
            for point in description_points:
                if point.strip():
                    self._add_paragraph(f"• {point.strip()}", spacing=1)

    def _add_education_section(self, education: List[Dict[str, Any]]) -> None:
        self._add_paragraph("Education", bold=True, add_border=True)
        for edu in education:
            edu_institution_abbr = "".join([word[0] for word in edu['edu_institution'].split(" ") if not word[0].islower()])
            self._add_two_column_table(
                f"{edu['edu_course']} | {edu_institution_abbr}",
                f"{self._format_date(edu['edu_start_date'])} - {self._format_date(edu['edu_end_date'])}",
                bold=False
            )

    def generate(self, data: Dict[str, Any], output_path: str) -> None:
        # Set the default font for the document
        style = self.document.styles['Normal']
        font = style.font
        font.name = "Times New Roman"  # Set default font to Times New Roman
        font.size = Pt(self.DEFAULT_FONT_SIZE)
        sections = [
            Section("Header", lambda: self._add_header_section(data)),
            Section("Experience", lambda: self._add_experience_section(data['experiences'])),
            Section("Skills", lambda: self._add_skills_section(data['user_skills'])),
            Section("Projects", lambda: self._add_projects_section(data['projects'])),
            Section("Education", lambda: self._add_education_section(data['education']))
        ]
        for section in sections:
            try:
                section.content_generator()
            except Exception as e:
                raise ValueError(f"Error generating {section.title} section: {str(e)}")
        style = self.document.styles['Normal']
        style.paragraph_format.line_spacing = 1
        self.document.save(output_path)
         # Convert DOCX to PDF
        try:
            convert(output_path)
        except Exception as e:
            raise ValueError(f"Error converting DOCX to PDF: {str(e)}")

