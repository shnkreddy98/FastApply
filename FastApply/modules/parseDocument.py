from docx import Document
import re
from datetime import datetime
import json

import streamlit as st

class parseDocument():

    def __init__(self):
        pass

    def get_datetime(date_str):
            try:
                return datetime.strptime(date_str, "%b-%Y").date()
            except ValueError:
                return None
            
    def put_datestr(d):
        return d.strftime("%b %Y") if d else "Present"
    
    def convert_datetime_in_string(json_string):
        # Pattern to match datetime.date(Y, M, D)
        date_pattern = re.compile(r"datetime\.date\((\d{4}),\s*(\d{1,2}),\s*(\d{1,2})\)")

        # Replace with "YYYY-MM-DD" format
        formatted_string = date_pattern.sub(r'"\1-\2-\3"', json_string)
        return formatted_string

    def read_file(documentfilename):
        """
        Input: Path to the directory in which resume resides or the Path to the docx file itself
        Output: Text version of the resume
        Function: Extracts data from word file
        """
        print(f"Reading file from {documentfilename}")
        document = Document(documentfilename)
        # text = ''
        # for paragraph in document.paragraphs:
        #     text += (f"\n {paragraph.text}")
        return document
    
    def save_file(document, filename="FastApply/output/shashankreddy.docx"):
        """
        Input: String value of Job description
        Output: 
        Function: Saves text/str data to word document
        """
        document.save(filename)
        print(f"Resume Saved at {filename}")
    
    def extract_details(docx_file):
        doc = Document(docx_file)
        data = {
            "name": "",
            "location": "",
            "email": None,
            "phone": None,
            "summary": "",
            "work_experience": [],
            "skills": "",
            "projects": [],
            "education": []
        }
        
        content = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Extract Name (First line assumed to be name)
        data["name"] = content[0] if content else ""
        
        # Extract Contact Details (Assuming it is in the second line)
        contact_line = content[1] if len(content) > 1 else ""
        location_match = re.search(r'^[^|]+', contact_line)
        data["location"] = location_match.group().strip() if location_match else None
        
        email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', contact_line)
        data["email"] = email_match.group() if email_match else None
        
        phone_match = re.search(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', contact_line)
        data["phone"] = phone_match.group() if phone_match else None
        
        # Section identifiers
        current_section = None
        for line in content[2:]:
            if line.lower().startswith("summary"):
                current_section = "summary"
            elif line.lower().startswith("work experience"):
                current_section = "work_experience"
            elif line.lower().startswith("skills"):
                current_section = "skills"
            elif line.lower().startswith("academic projects"):
                current_section = "projects"
            elif line.lower().startswith("education"):
                current_section = "education"
            elif current_section == "summary":
                data["summary"] += line + " "
            elif current_section == "work_experience":
                match = re.match(r'^(.*?)\s*\|\s*(.*?)\s*\|\s*.*?(\w{3}\s\d{4})\s+[–-]\s+(\w{3}\s\d{4})$', line)
                if match:
                    company, role, start_date, end_date = match.groups()
                    data["work_experience"].append({
                        "work_exp_company": company.strip(),
                        "work_exp_role": role.strip(),
                        "work_exp_start_date": parseDocument.get_datetime(start_date),
                        "work_exp_end_date": parseDocument.get_datetime(end_date),
                        "work_exp_description": "",
                        "work_exp_location": ""
                    })
                else:
                    if data["work_experience"]:
                        data["work_experience"][-1]["work_exp_description"] += line + " "
            elif current_section == "skills":
                data["skills"] += line + " "
            elif current_section == "projects":
                match = re.match(r'^(.*?)\s*\|\s*.*?(\w{3}\s\d{4})\s+[–-]\s+(\w{3}\s\d{4})$', line)
                if match:
                    project_name, start_date, end_date = match.groups()
                    data["projects"].append({
                        "proj_name": project_name.strip(),
                        "proj_start_date": parseDocument.get_datetime(start_date),
                        "proj_end_date": parseDocument.get_datetime(end_date),
                        "proj_description": ""
                    })
                else:
                    if data["projects"]:
                        data["projects"][-1]["proj_description"] += line + " "
            elif current_section == "education":
                match = re.match(r'^(.*?)\s*\|\s*(.*?)\s+(\w{3}\s\d{4})\s+[–-]\s+(\w{3}\s\d{4})$', line)
                if match:
                    course, institution, start_date, end_date = match.groups()
                    data["education"].append({
                        "edu_course": course.strip(),
                        "edu_institution": institution.strip(),
                        "edu_start_date": parseDocument.get_datetime(start_date),
                        "edu_end_date": parseDocument.get_datetime(end_date)
                    })
        
        # Clean up trailing spaces
        data["summary"] = data["summary"].strip()
        data["skills"] = data["skills"].strip()
        return data

        