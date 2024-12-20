from docx import Document

import streamlit as st

class structureDocx():

    def __init__(self):
        pass

    def create_document(new_resume_text, user_name):
        # Load the existing document
        resume_template = Document("FastApply/inputs/template.docx")
        new_resume = Document()

        new_resume_text = new_resume_text.split("\n")

        # Iterate through paragraphs and add the new text into the new document
        for i, paragraph in enumerate(resume_template.paragraphs):
            if i < len(new_resume_text):
                new_paragraph = new_resume.add_paragraph(new_resume_text[i])
                
                # Copying formatting
                for run in paragraph.runs:
                    new_paragraph.add_run(new_resume_text[i]).bold = run.bold
                    new_paragraph.add_run(new_resume_text[i]).italic = run.italic
                    new_paragraph.add_run(new_resume_text[i]).underline = run.underline
                    new_paragraph.add_run(new_resume_text[i]).font.size = run.font.size
                    new_paragraph.add_run(new_resume_text[i]).font.name = run.font.name

        # Save the new document
        resume_template.save(f"FastApply/outputs/{user_name}.docx")

