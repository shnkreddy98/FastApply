from docx import Document
import re
from datetime import datetime

class parseDocument():

    def __init__(self):
        pass

    def format_date(date_str):
            try:
                return datetime.strptime(date_str, "%b-%Y").date()
            except ValueError:
                return None

    def read_file(documentfilename):
        """
        Input: Path to the directory in which resume resides or the Path to the docx file itself
        Output: Text version of the resume
        Function: Extracts data from word file
        """
        print(f"Reading file from {documentfilename}")
        document = Document(documentfilename)
        # text = ''
        # for paragraph in document.paragraphs:
        #     text += (f"\n {paragraph.text}")
        return document
    
    def save_file(document, filename="ResumeBuilder/output/shashankreddy.docx"):
        """
        Input: String value of Job description
        Output: 
        Function: Saves text/str data to word document
        """
        document.save(filename)
        print(f"Resume Saved at {filename}")

    def extract_details(docx_file):
        doc = Document(docx_file)
        data = {
            "name": "",
            "location": "",
            "email": "",
            "phone": "",
            "summary": "",
            "work_experience": [],
            "skills": "",
            "projects": [],
            "education": []
        }
        
        content = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Extract Name (First line assumed to be name)
        data["name"] = content[0]
        
        # Extract Contact Details (Assuming it is in the second line)
        contact_line = content[1]
        data["location"] = re.search(r'^[^|]+', contact_line).group().strip()
        data["email"] = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', contact_line).group()
        data["phone"] = re.search(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', contact_line).group()
        
        # Section identifiers
        current_section = None
        for line in content[2:]:
            if line.lower().startswith("summary"):
                current_section = "summary"
            elif line.lower().startswith("work experience"):
                current_section = "work_experience"
            elif line.lower().startswith("skills"):
                current_section = "skills"
            elif line.lower().startswith("academic projects"):
                current_section = "projects"
            elif line.lower().startswith("education"):
                current_section = "education"
            elif current_section == "summary":
                data["summary"] += line + " "
            elif current_section == "work_experience":
                # Extract company, role, and dates (MMM-YYYY format)
                match = re.match(r'^(.*?)\s*\|\s*(.*?)\s*\|\s*.*?(\w{3}\s\d{4})\s+[–-]\s+(\w{3}\s\d{4})$', line)
                if match:
                    company, role, start_date, end_date = match.groups()
                    data["work_experience"].append({
                        "company": company.strip(),
                        "role": role.strip(),
                        "start_date": parseDocument.format_date(start_date),
                        "end_date": parseDocument.format_date(end_date),
                        "description": ""
                    })
                else:
                    if data["work_experience"]:
                            data["work_experience"][-1]["description"] = data["work_experience"][-1].get("description", "") + line + " "
            elif current_section == "skills":
                data["skills"] += line + " "
            elif current_section == "projects":
                # Extract project details (MMM-YYYY format)
                match = re.match(r'^(.*?)\s*\|\s*.*?(\w{3}\s\d{4})\s+[–-]\s+(\w{3}\s\d{4})$', line)
                if match:
                    project_name, start_date, end_date = match.groups()
                    data["projects"].append({
                        "project_name": project_name.strip(),
                        "start_date": parseDocument.format_date(start_date),
                        "end_date": parseDocument.format_date(end_date),
                        "description": ""
                    })
                else:
                    if data["projects"]:
                        data["projects"][-1]["description"] += line + " "
            elif current_section == "education":
                # Extract education details (MMM-YYYY format)
                match = re.match(r'^(.*?)\s*\|\s*(.*?)\s+(\w{3}\s\d{4})\s+[–-]\s+(\w{3}\s\d{4})$', line)
                if match:
                    course, institution, start_date, end_date = match.groups()
                    data["education"].append({
                        "course": course.strip(),
                        "institution": institution.strip(),
                        "start_date": parseDocument.format_date(start_date),
                        "end_date": parseDocument.format_date(end_date)
                    })
        
        # Clean up trailing spaces
        data["summary"] = data["summary"].strip()
        data["skills"] = data["skills"].strip()
        return data

